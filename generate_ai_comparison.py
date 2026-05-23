"""
Generate AI Tools Comparison Word document
Output: ai-tools-comparison.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\Users\alons\hebrew-quiz-bot\ai-tools-comparison.docx"

# Colors
GREEN_DARK  = RGBColor(0x2e, 0x7d, 0x32)
GREEN_MID   = RGBColor(0x43, 0xa0, 0x47)
GREEN_LIGHT = RGBColor(0xe8, 0xf5, 0xe9)
WHITE       = RGBColor(0xff, 0xff, 0xff)
GREY        = RGBColor(0x44, 0x44, 0x44)
GREY_LIGHT  = RGBColor(0x77, 0x77, 0x77)
GOLD        = RGBColor(0xff, 0x8f, 0x00)
BLUE_DARK   = RGBColor(0x0d, 0x47, 0xa1)
RED_DARK    = RGBColor(0xb7, 0x1c, 0x1c)
ORANGE      = RGBColor(0xe6, 0x51, 0x00)
TEAL        = RGBColor(0x00, 0x69, 0x5c)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color='C8E6C9', sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top','bottom','left','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)

def set_table_border(table, color='DDDDDD'):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ('top','bottom','left','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)

def para_border_bottom(para, color='43A047'):
    pPr = para._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pb.append(bottom)
    pPr.append(pb)

def set_rtl(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_heading(doc, text, level=1, rtl=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(8)
    if rtl:
        set_rtl(p)
    r = p.add_run(text)
    r.font.bold = True
    r.font.name = 'Calibri'
    if level == 1:
        r.font.size = Pt(20)
        r.font.color.rgb = GREEN_DARK
        para_border_bottom(p)
    elif level == 2:
        r.font.size = Pt(14)
        r.font.color.rgb = GREEN_MID
    else:
        r.font.size = Pt(12)
        r.font.color.rgb = GREY
    return p

def add_para(doc, text, rtl=True, size=10.5, color=None, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(space_after)
    if rtl:
        set_rtl(p)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.font.color.rgb = color or GREY
    r.font.bold = bold
    return p

def stars(n, total=5):
    return '★' * n + '☆' * (total - n)

# ── Data ─────────────────────────────────────────────────────────────────────

TOOLS = ["Claude\n(Anthropic)", "ChatGPT-4o\n(OpenAI)", "GitHub\nCopilot", "Gemini\nAdvanced", "Cursor AI"]
TOOL_COLORS = ['2E7D32', '0D47A1', '37474F', '1565C0', '4A148C']
TOOL_NOTES  = ["✓ נבחר", "", "", "", ""]

CRITERIA = [
    # (קריטריון, תיאור קצר, [ציונים 1-5], [הערות לכל כלי])
    (
        "תמיכה בעברית ו-RTL",
        "יכולת עיבוד כתיבה מימין לשמאל, הבנת הקשר תרבותי ישראלי",
        [5, 4, 2, 4, 2],
        ["הבנה מעולה של RTL, בעיות ניואנס עברי", "תמיכה טובה, לעיתים שגיאות בהקשר תרבותי", "מוגבל — ממוקד בקוד בלבד", "תמיכה טובה בזכות Google Translate", "תלוי במודל הבסיסי — לרוב GPT"],
    ),
    (
        "הבנת הקשר המלא של הפרויקט",
        "יכולת לזכור ולהשתמש בהקשר רב-קבצים לאורך שיחה ארוכה",
        [5, 3, 2, 3, 4],
        ["חלון הקשר גדול מאוד — זוכר את כל השיחה", "חלון הקשר מוגבל יותר, נוטה 'לשכוח'", "רואה רק קובץ אחד בכל פעם", "חלון גדול אך פחות עקבי", "אינטגרציה עם קוד קיים — טוב מאוד"],
    ),
    (
        "יצירת קוד React ו-Firebase",
        "איכות קוד מודרני: hooks, Firestore, Auth, Realtime",
        [5, 5, 4, 4, 5],
        ["קוד נקי, תבניות מודרניות, הבנת Firebase מעמיקה", "קוד מצוין, לפעמים דפוסים מיושנים", "השלמה טובה, לא מסביר או מנחה", "קוד טוב, לפעמים verbose", "ייעוץ ועריכה ישירים בסביבת הפיתוח"],
    ),
    (
        "עיצוב UI/UX",
        "הצעות לממשק משתמש, CSS, נגישות ומובייל",
        [5, 3, 1, 3, 2],
        ["הצעות מפורטות לעיצוב, glass morphism, RTL נכון", "מסוגל, אך לא חזק בעיצוב ויזואלי", "לא רלוונטי — כלי קוד בלבד", "בסיסי, ללא עומק עיצובי", "מוגבל — ממוקד בפרודוקטיביות קוד"],
    ),
    (
        "תיעוד ומסמכים",
        "יצירת README, GDD, user stories, Use Cases",
        [5, 4, 1, 3, 1],
        ["יצר GDD, 10 מסמכי עיצוב, 100 user stories, Word ו-PDF", "מסוגל, דורש הנחיות מפורטות", "לא מיועד לתיעוד כלל", "בסיסי, פחות מובנה", "לא מיועד לתיעוד כלל"],
    ),
    (
        "זיהוי ותיקון באגים",
        "יכולת למצוא סיבת שורש של בעיות ולתקן בצורה מדויקת",
        [5, 4, 3, 3, 4],
        ["מאתר סיבת שורש, מסביר ומתקן — כולל בעיות RTL עדינות", "טוב, לפעמים מציע פתרונות עוקפים", "מגיש הצעות, לא תמיד מסביר למה", "סביר, פחות מדויק בבאגים מורכבים", "טוב בזכות גישה ישירה לקוד"],
    ),
    (
        "חינם / מחיר",
        "עלות שימוש לפרויקט אקדמי",
        [4, 3, 2, 3, 3],
        ["Claude.ai free tier זמין; Pro $20/חודש", "Free tier מוגבל; Plus $20/חודש", "דורש מנוי GitHub/VS Code — $10-19/חודש", "Gemini Advanced כלול ב-Google One $20/חודש", "Free tier; Pro $20/חודש"],
    ),
    (
        "אינטגרציה עם כלי פיתוח",
        "שילוב עם VS Code, Terminal, GitHub",
        [5, 3, 5, 2, 5],
        ["Claude Code CLI — ישירות בטרמינל ו-VS Code", "דרך API או אתר בלבד, אין CLI ייעודי", "שילוב מלא ב-VS Code ו-JetBrains", "אין אינטגרציה ייעודית לסביבת פיתוח", "IDE מובנה עם Claude/GPT"],
    ),
    (
        "עקביות לאורך זמן",
        "שמירת החלטות עיצוב ותכנות לאורך הפרויקט",
        [5, 3, 2, 3, 3],
        ["זיכרון פרויקט, הנחיות בקובץ CLAUDE.md", "נוטה לסתור החלטות קודמות", "אין מנגנון זיכרון — כל שלמה עצמאית", "עקביות בינונית", "עקבי בתוך session, אין זיכרון בין סשנים"],
    ),
    (
        "הסברים פדגוגיים",
        "יכולת להסביר קוד ומושגים לסטודנט לומד",
        [5, 4, 1, 4, 2],
        ["מסביר למה, לא רק מה — מותאם לרמת הלומד", "הסברים טובים, לפעמים ארוכים מדי", "לא מסביר כלל — רק משלים", "הסברים נאים אך גנריים", "ממוקד בפרודוקטיביות, לא בלמידה"],
    ),
]

SCORE_COLORS = {
    5: ('1B5E20', 'E8F5E9'),  # dark green bg, light green fill
    4: ('2E7D32', 'F1F8E9'),
    3: ('F57F17', 'FFFDE7'),
    2: ('E65100', 'FFF3E0'),
    1: ('B71C1C', 'FFEBEE'),
}

# ── Document ──────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
    doc.styles['Normal'].font.name = 'Calibri'

    # ── Cover ─────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '2E7D32')
    pPr.append(shd)
    r = p.add_run('  '); r.font.size = Pt(32)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    r = p.add_run('השוואת כלי AI לפיתוח פרויקט')
    r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GREEN_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('תחום הפונקציה — משחק למידה בעברית')
    r.font.size = Pt(16); r.font.color.rgb = GREEN_MID

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('─' * 40)
    r.font.color.rgb = GREEN_MID; r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run('אלון שובה ויוסי לוין  ·  מכללת סמי שמעון  ·  2026')
    r.font.size = Pt(11); r.font.color.rgb = GREY_LIGHT

    doc.add_page_break()

    # ── Introduction ──────────────────────────────────────────────────────────
    add_heading(doc, 'מבוא', 1)
    add_para(doc,
        'במסגרת פרויקט הגמר "תחום הפונקציה", נדרשנו לבחור כלי AI שילווה אותנו לאורך כל שלבי הפיתוח — '
        'מתכנון הארכיטקטורה ועד לעיצוב הממשק, כתיבת הקוד, איתור באגים ויצירת תיעוד מקצועי. '
        'מסמך זה מציג השוואה מקיפה בין חמישה כלי AI מובילים, ומנמק מדוע הבחירה ב-Claude של Anthropic '
        'היוותה את ההחלטה האסטרטגית הנכונה ביותר עבור פרויקט זה.',
        rtl=True, size=11, space_after=8)

    add_para(doc,
        'הכלים שנבדקו: Claude (Anthropic), ChatGPT-4o (OpenAI), GitHub Copilot (Microsoft), '
        'Gemini Advanced (Google), ו-Cursor AI.',
        rtl=True, size=11, space_after=4)

    add_para(doc,
        'הקריטריונים נבחרו בהתאם לצרכי הפרויקט הספציפיים: תמיכה בעברית ו-RTL, '
        'פיתוח React ו-Firebase, עיצוב UI/UX, יצירת תיעוד, ועקביות לאורך פרויקט ממושך.',
        rtl=True, size=11, space_after=16)

    # ── Main comparison table ─────────────────────────────────────────────────
    add_heading(doc, 'טבלת השוואה מרכזית', 1)

    # Header note
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run('ציון: ★★★★★ = מעולה  |  ★★★★☆ = טוב מאוד  |  ★★★☆☆ = סביר  |  ★★☆☆☆ = חלש  |  ★☆☆☆☆ = לא מתאים')
    r.font.size = Pt(9); r.font.color.rgb = GREY_LIGHT; r.font.italic = True

    cols = 1 + len(TOOLS)
    table = doc.add_table(rows=1 + len(CRITERIA), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table, 'CCCCCC')

    # Set column widths
    widths = [Cm(5.2)] + [Cm(3.0)] * len(TOOLS)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    # Header row
    header_row = table.rows[0]
    set_cell_bg(header_row.cells[0], '1B5E20')
    h = header_row.cells[0].paragraphs[0]
    h.paragraph_format.space_before = Pt(4); h.paragraph_format.space_after = Pt(4)
    set_rtl(h)
    rr = h.add_run('קריטריון')
    rr.font.bold = True; rr.font.size = Pt(10); rr.font.color.rgb = WHITE

    for i, (tool, color, note) in enumerate(zip(TOOLS, TOOL_COLORS, TOOL_NOTES)):
        cell = header_row.cells[i + 1]
        set_cell_bg(cell, color)
        set_cell_border(cell, 'FFFFFF', 6)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(tool)
        r.font.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = WHITE
        if note:
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(note)
            r2.font.size = Pt(8); r2.font.color.rgb = WHITE; r2.font.bold = True

    # Data rows
    for ri, (crit, desc, scores, notes) in enumerate(CRITERIA):
        row = table.rows[ri + 1]
        bg = 'F9FBE7' if ri % 2 == 0 else 'FFFFFF'

        # Criterion cell
        cell0 = row.cells[0]
        set_cell_bg(cell0, bg)
        set_cell_border(cell0, 'DDDDDD')
        p = cell0.paragraphs[0]
        set_rtl(p)
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
        rr = p.add_run(crit)
        rr.font.bold = True; rr.font.size = Pt(9.5); rr.font.color.rgb = GREEN_DARK
        p2 = cell0.add_paragraph()
        set_rtl(p2)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(desc)
        r2.font.size = Pt(8); r2.font.color.rgb = GREY_LIGHT; r2.font.italic = True

        # Score cells
        for ci, score in enumerate(scores):
            cell = row.cells[ci + 1]
            fg, fill = SCORE_COLORS[score]
            set_cell_bg(cell, fill)
            set_cell_border(cell, 'DDDDDD')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(stars(score))
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(
                int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16)
            )
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(4)
            r2 = p2.add_run(f'{score}/5')
            r2.font.size = Pt(8); r2.font.bold = True
            r2.font.color.rgb = RGBColor(int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16))

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ── Total scores ──────────────────────────────────────────────────────────
    add_heading(doc, 'סיכום ציונים', 2)

    totals = [sum(c[2][i] for c in CRITERIA) for i in range(len(TOOLS))]
    max_score = len(CRITERIA) * 5

    score_table = doc.add_table(rows=2, cols=len(TOOLS))
    score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(score_table, 'CCCCCC')

    tool_names_short = ["Claude", "ChatGPT-4o", "GitHub Copilot", "Gemini", "Cursor AI"]
    for i, (name, color, total) in enumerate(zip(tool_names_short, TOOL_COLORS, totals)):
        # Name row
        c1 = score_table.rows[0].cells[i]
        set_cell_bg(c1, color)
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(5)
        r = p.add_run(name)
        r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE

        # Score row
        c2 = score_table.rows[1].cells[i]
        fill = 'E8F5E9' if i == 0 else 'FFFFFF'
        set_cell_bg(c2, fill)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(5); p2.paragraph_format.space_after = Pt(5)
        r2 = p2.add_run(f'{total}/{max_score}')
        r2.font.bold = True
        r2.font.size = Pt(14 if i == 0 else 11)
        r2.font.color.rgb = GREEN_DARK if i == 0 else GREY

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── Detailed analysis per tool ────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, 'ניתוח מפורט לכל כלי', 1)

    TOOL_DETAIL = [
        (
            "Claude (Anthropic) — הכלי שנבחר ✓",
            '2E7D32',
            [
                "חלון הקשר של עד 200,000 טוקנים — מאפשר לנהל שיחה ממושכת על כל הפרויקט ללא אובדן הקשר.",
                "Claude Code CLI — אינטגרציה ישירה עם הטרמינל ו-VS Code, כולל עריכת קבצים, הרצת builds ו-deploy.",
                "הבנה עמוקה של עברית ו-RTL: זיהה ותיקן בעיות עיצוב ייחודיות לכיוון הכתיבה (כפתורים בצד הלא נכון, פונקציות שמופיעות הפוך).",
                "קובץ CLAUDE.md — מנגנון זיכרון פרויקטלי שמאפשר עקביות בין סשנים שונים.",
                "יצר תיעוד מקיף: GDD, 10 מסמכי עיצוב, 100 user stories, Use Cases, קובצי Word ו-PDF.",
                "הסברים פדגוגיים מותאמים — מסביר למה ולא רק מה, מה שסייע ללמידה מקבילה לפיתוח.",
            ],
            "הניקוד הגבוה ביותר בכל הקריטריונים שחשובים לפרויקט זה."
        ),
        (
            "ChatGPT-4o (OpenAI)",
            '0D47A1',
            [
                "מודל חזק ביצירת קוד React ו-Firebase, ידע רחב ועדכני.",
                "ממשק web נוח, אך ללא CLI ייעודי לפיתוח — דורש העתקה ידנית של קוד.",
                "חלון הקשר קטן יותר — בפרויקטים ממושכים נוטה 'לשכוח' החלטות קודמות.",
                "תמיכה בעברית סבירה, אך נתקל בבעיות ב-RTL עדין (כיוון כפתורים, סדר אלמנטים).",
                "אין מנגנון זיכרון בין שיחות — כל session מתחיל מאפס.",
            ],
            "כלי מצוין לשאלות נקודתיות, אך פחות מתאים לפרויקט ארוך עם המשכיות."
        ),
        (
            "GitHub Copilot (Microsoft)",
            '37474F',
            [
                "השלמת קוד חזקה בתוך ה-IDE, מהירה ולא מפריעה לזרימת העבודה.",
                "לא מסביר קוד, לא מנחה, לא יוצר תיעוד — כלי השלמה בלבד.",
                "רואה קובץ אחד בכל פעם, ללא הבנת ארכיטקטורה רחבה.",
                "לא מתאים לתמיכה בעברית, עיצוב UI, או יצירת מסמכים.",
                "מחיר גבוה יחסית לערך שמספק בפרויקט אקדמי.",
            ],
            "שימושי כתוסף, לא כלי ראשי לפרויקט מורכב."
        ),
        (
            "Gemini Advanced (Google)",
            '1565C0',
            [
                "חלון הקשר גדול (1 מיליון טוקנים בגרסאות מסוימות), אך פחות עקבי בשימוש בו.",
                "אינטגרציה עם Google Workspace — יתרון לכתיבת מסמכים ב-Google Docs.",
                "קוד סביר, אך לפעמים verbose ומיושן בהשוואה לאלטרנטיבות.",
                "אין CLI ייעודי לפיתוח, אין אינטגרציה ישירה עם הטרמינל.",
                "תמיכה בעברית טובה בזכות Google Translate, אך חסרה הבנה תרבותית עמוקה.",
            ],
            "כלי מבטיח לעתיד, אך כיום פחות בשל מ-Claude לפרויקטי קוד."
        ),
        (
            "Cursor AI",
            '4A148C',
            [
                "IDE מובנה עם AI — חווית פיתוח חלקה עם עריכה ישירה בקוד.",
                "תומך ב-Claude ו-GPT-4 כמנועים, כך שאיכות הקוד תלויה בבחירת המנוע.",
                "חזק בעריכה מולטי-קובץ בתוך ה-IDE, אך פחות גמיש מ-CLI.",
                "אין מנגנון תיעוד, memory, או זיכרון פרויקטלי מובנה.",
                "לא תומך בשיחה מחוץ לסביבת הקוד — פחות שימושי לתכנון ועיצוב.",
            ],
            "אלטרנטיבה מצוינת לעריכת קוד, אך לא מחליף את Claude לניהול פרויקט מלא."
        ),
    ]

    for name, color, points, summary in TOOL_DETAIL:
        # Tool banner
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color)
        pPr.append(shd)
        set_rtl(p)
        r = p.add_run(f'  {name}  ')
        r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = WHITE

        for pt in points:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Inches(0.25)
            set_rtl(p)
            rr = p.add_run('◆  ')
            rr.font.size = Pt(7)
            rr.font.color.rgb = RGBColor(int(color[0:2],16), int(color[2:4],16), int(color[4:6],16))
            r2 = p.add_run(pt)
            r2.font.size = Pt(10); r2.font.color.rgb = GREY

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(10)
        p.paragraph_format.left_indent  = Inches(0.25)
        set_rtl(p)
        pPr = p._p.get_or_add_pPr()
        shd2 = OxmlElement('w:shd')
        shd2.set(qn('w:val'), 'clear'); shd2.set(qn('w:color'), 'auto')
        shd2.set(qn('w:fill'), 'F5F5F5')
        pPr.append(shd2)
        r = p.add_run('📌  ')
        r.font.size = Pt(10)
        r2 = p.add_run(summary)
        r2.font.size = Pt(10); r2.font.italic = True; r2.font.color.rgb = GREY

    # ── Conclusion ───────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, 'מסקנה — מדוע Claude היה הבחירה הנכונה', 1)

    conclusions = [
        (
            "1. הפרויקט דרש הבנת עברית ו-RTL ברמה גבוהה",
            "הפרויקט כולו בעברית — ממשק, שאלות, כפתורים, ופעילות RTL עדינה (סדר אלמנטים, כיוון אנימציות, "
            "כיוון פונקציות מתמטיות צפות). Claude זיהה ותיקן בעיות RTL שכלים אחרים כלל לא ידעו שקיימות."
        ),
        (
            "2. פרויקט ממושך — נדרשת עקביות לאורך זמן",
            "פרויקט גמר הוא מרתון, לא ספרינט. CLAUDE.md ומנגנון הזיכרון הבטיחו שהחלטות עיצוב ותכנות "
            "נשמרו לאורך כל שלבי הפיתוח, מבלי לחזור על הסברים בכל שיחה מחדש."
        ),
        (
            "3. Claude Code CLI — שינה את כללי המשחק",
            "היכולת לעבוד ישירות בטרמינל ו-VS Code — לערוך קבצים, להריץ builds, לעשות deploy ל-Firebase, "
            "ולקרוא שגיאות — הפכה את הפיתוח ליעיל פי כמה לעומת כלי web בלבד."
        ),
        (
            "4. תיעוד מקצועי — לא רק קוד",
            "פרויקט גמר דורש תיעוד: GDD, Use Cases, User Stories, ומסמכי עיצוב. Claude יצר את כל אלה "
            "בעברית תקנית, בפורמט Word ו-PDF מקצועי — יכולת שאין לה תחרות בקרב הכלים שנבדקו."
        ),
        (
            "5. עלות-תועלת מיטבית לפרויקט אקדמי",
            "Claude.ai free tier סיפק את רוב הצרכים. Claude Code CLI זמין ב-Max plan — השקעה שהחזירה "
            "את עצמה בחיסכון הזמן העצום שבפיתוח, debugging, ויצירת תיעוד."
        ),
    ]

    for title, body in conclusions:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        set_rtl(p)
        r = p.add_run(title)
        r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = GREEN_DARK

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(8)
        p2.paragraph_format.left_indent = Inches(0.2)
        set_rtl(p2)
        pPr = p2._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F9FBE7')
        pPr.append(shd)
        r2 = p2.add_run(body)
        r2.font.size = Pt(11); r2.font.color.rgb = GREY

    # Final verdict box
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1B5E20')
    pPr.append(shd)
    set_rtl(p)
    r = p.add_run('  ✅  סיכום סופי  ')
    r.font.bold = True; r.font.size = Pt(14); r.font.color.rgb = WHITE

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_rtl(p2)
    pPr2 = p2._p.get_or_add_pPr()
    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:val'), 'clear'); shd2.set(qn('w:color'), 'auto'); shd2.set(qn('w:fill'), 'E8F5E9')
    pPr2.append(shd2)
    r2 = p2.add_run(
        '  מתוך חמשת הכלים שנבדקו, Claude של Anthropic השיג את הציון הגבוה ביותר — '
        f'{totals[0]}/{max_score} — ובלט בכל הפרמטרים הקריטיים לפרויקט זה: '
        'עברית ו-RTL, אינטגרציה ב-CLI, עקביות, תיעוד ופדגוגיה. '
        'הבחירה בו לא הייתה מקרית — היא הייתה ההחלטה המקצועית הנכונה.  '
    )
    r2.font.size = Pt(11.5); r2.font.color.rgb = GREEN_DARK; r2.font.bold = True

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

if __name__ == "__main__":
    build()
