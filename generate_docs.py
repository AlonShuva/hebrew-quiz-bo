"""
Generate a Word document from all markdown files in docs/
Outputs: project-documentation.docx
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DOCS_DIR = r"c:\Users\alons\hebrew-quiz-bot\docs"
OUTPUT    = r"c:\Users\alons\hebrew-quiz-bot\project-documentation.docx"

GREEN_DARK   = RGBColor(0x2e, 0x7d, 0x32)
GREEN_MID    = RGBColor(0x43, 0xa0, 0x47)
GREEN_LIGHT  = RGBColor(0xe8, 0xf5, 0xe9)
WHITE        = RGBColor(0xff, 0xff, 0xff)
GREY_DARK    = RGBColor(0x33, 0x33, 0x33)
GREY_LIGHT   = RGBColor(0x66, 0x66, 0x66)
CODE_BG      = RGBColor(0xf5, 0xf5, 0xf5)
CODE_FG      = RGBColor(0x1a, 0x23, 0x7e)

FILE_ORDER = [
    "game-design-document.md",
    "user-flow.md",
    "economy-points-system.md",
    "multiplayer-system.md",
    "firebase-architecture.md",
    "firestore-collections.md",
    "admin-panel.md",
    "achievement-system.md",
    "leaderboards.md",
    "adaptive-difficulty.md",
]

# ── helpers ──────────────────────────────────────────────────────────────────

def set_rtl(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)

def set_cell_border(cell, sides=('top','bottom','left','right'), color='C8E6C9', sz=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)

def set_para_border_bottom(para, color='C8E6C9'):
    pPr = para._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pb.append(bottom)
    pPr.append(pb)

def set_table_border(table, color='C8E6C9'):
    tbl    = table._tbl
    tblPr  = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ('top','bottom','left','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)

def inline_style(run, text):
    """Apply bold/italic/code styling from inline markdown."""
    # strip bold+italic
    if text.startswith('***') and text.endswith('***'):
        run.bold   = True
        run.italic = True
        return text[3:-3]
    if text.startswith('**') and text.endswith('**'):
        run.bold = True
        return text[2:-2]
    if text.startswith('*') and text.endswith('*'):
        run.italic = True
        return text[1:-1]
    if text.startswith('`') and text.endswith('`'):
        run.font.name  = 'Courier New'
        run.font.size  = Pt(9)
        run.font.color.rgb = CODE_FG
        return text[1:-1]
    return text

def add_inline(para, text, base_size=10, base_color=GREY_DARK, bold=False):
    """Parse inline markdown (bold/italic/code) and add runs to para."""
    pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)'
    parts   = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        run = para.add_run()
        run.font.size  = Pt(base_size)
        run.font.color.rgb = base_color
        run.bold       = bold
        clean = inline_style(run, part)
        run.text = clean

# ── document setup ────────────────────────────────────────────────────────────

def setup_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # Default style
    style = doc.styles['Normal']
    style.font.name  = 'Calibri'
    style.font.size  = Pt(10)
    style.font.color.rgb = GREY_DARK

    return doc

# ── cover page ────────────────────────────────────────────────────────────────

def add_cover(doc):
    # Top color band
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '2E7D32')
    pPr.append(shd)
    run = p.add_run('  ')
    run.font.size = Pt(36)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run('תחום הפונקציה')
    r.font.name  = 'Calibri'
    r.font.size  = Pt(36)
    r.font.bold  = True
    r.font.color.rgb = GREEN_DARK

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('Hebrew Math Learning Game')
    r.font.size  = Pt(22)
    r.font.color.rgb = GREEN_MID

    # Divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('─' * 38)
    r.font.color.rgb = GREEN_MID
    r.font.size = Pt(12)

    # Description
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run('Project Documentation')
    r.font.size  = Pt(14)
    r.font.color.rgb = GREY_LIGHT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('React · Firebase · Vite')
    r.font.size  = Pt(11)
    r.font.color.rgb = GREY_LIGHT

    # Version / date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run('Version 1.0  ·  2026')
    r.font.size  = Pt(10)
    r.font.color.rgb = GREY_LIGHT

    doc.add_page_break()

# ── table of contents ─────────────────────────────────────────────────────────

DOC_TITLES = {
    "game-design-document.md":    "Game Design Document",
    "user-flow.md":               "User Flow",
    "economy-points-system.md":   "Economy & Points System",
    "multiplayer-system.md":      "Multiplayer System Design",
    "firebase-architecture.md":   "Firebase Architecture",
    "firestore-collections.md":   "Firestore Collections",
    "admin-panel.md":             "Admin Panel",
    "achievement-system.md":      "Achievement System",
    "leaderboards.md":            "Leaderboard System",
    "adaptive-difficulty.md":     "Adaptive Difficulty System",
}

def add_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run('Table of Contents')
    r.font.size  = Pt(20)
    r.font.bold  = True
    r.font.color.rgb = GREEN_DARK
    set_para_border_bottom(p)
    p.paragraph_format.space_after = Pt(16)

    for i, fname in enumerate(FILE_ORDER, 1):
        title = DOC_TITLES.get(fname, fname)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(f'  {i:02d}.  {title}')
        r.font.size  = Pt(11)
        r.font.color.rgb = GREY_DARK

    doc.add_page_break()

# ── markdown renderer ─────────────────────────────────────────────────────────

def render_markdown(doc, md_text, file_title):
    """Render a markdown string into the Word document."""

    # Section title (green banner)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(14)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '2E7D32')
    pPr.append(shd)
    r = p.add_run(f'  {file_title}  ')
    r.font.size  = Pt(18)
    r.font.bold  = True
    r.font.color.rgb = WHITE

    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # ── code block ──────────────────────────────────────────────────────
        if line.startswith('```'):
            if not in_code:
                in_code    = True
                code_lines = []
            else:
                in_code = False
                # render code block
                for cl in code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent  = Inches(0.3)
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after  = Pt(1)
                    pPr = p._p.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'),   'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'),  'F5F5F5')
                    pPr.append(shd)
                    r = p.add_run(cl)
                    r.font.name  = 'Courier New'
                    r.font.size  = Pt(8.5)
                    r.font.color.rgb = CODE_FG
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── horizontal rule ─────────────────────────────────────────────────
        if re.match(r'^-{3,}$', line.strip()):
            p = doc.add_paragraph()
            set_para_border_bottom(p, '43A047')
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            i += 1
            continue

        # ── headings ────────────────────────────────────────────────────────
        h_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if h_match:
            level = len(h_match.group(1))
            text  = h_match.group(2)
            # strip bold markers that are sometimes in headings
            text  = re.sub(r'\*\*(.*?)\*\*', r'\1', text)

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt([0, 18, 12, 8, 6][level])
            p.paragraph_format.space_after  = Pt([0, 8,  6,  4, 4][level])

            sizes  = [0, 18, 14, 12, 11]
            colors = [None, GREEN_DARK, GREEN_MID, GREEN_MID, GREY_DARK]

            r = p.add_run(text)
            r.font.size  = Pt(sizes[level])
            r.font.bold  = True
            r.font.color.rgb = colors[level]

            if level == 2:
                set_para_border_bottom(p)

            i += 1
            continue

        # ── table ────────────────────────────────────────────────────────────
        if line.startswith('|'):
            # collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1

            # filter separator row
            rows = [l for l in table_lines if not re.match(r'^\|[-| :]+\|$', l)]
            if not rows:
                continue

            # parse cells
            parsed = []
            for row in rows:
                cells = [c.strip() for c in row.strip('|').split('|')]
                parsed.append(cells)

            col_count = max(len(r) for r in parsed)
            table = doc.add_table(rows=len(parsed), cols=col_count)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            set_table_border(table)

            for ri, row in enumerate(parsed):
                for ci, cell_text in enumerate(row):
                    if ci >= col_count:
                        break
                    cell = table.cell(ri, ci)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    # strip inline markdown bold/code
                    clean = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                    clean = re.sub(r'`(.*?)`', r'\1', clean)

                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after  = Pt(3)
                    r = p.add_run(clean)
                    r.font.size = Pt(9)

                    if ri == 0:  # header row
                        r.font.bold       = True
                        r.font.color.rgb  = WHITE
                        set_cell_bg(cell, GREEN_DARK)
                    else:
                        r.font.color.rgb  = GREY_DARK
                        bg = GREEN_LIGHT if ri % 2 == 0 else WHITE
                        set_cell_bg(cell, bg)

                    set_cell_border(cell)

            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        # ── bullet list ──────────────────────────────────────────────────────
        bullet_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text   = bullet_match.group(3)

            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent  = Inches(0.2 + indent * 0.15)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.clear()
            add_inline(p, text, base_size=10)
            i += 1
            continue

        # ── blank line ───────────────────────────────────────────────────────
        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            i += 1
            continue

        # ── normal paragraph ─────────────────────────────────────────────────
        # skip the first H1 since we already rendered the banner
        if re.match(r'^# ', line):
            i += 1
            continue

        # skip metadata lines like **Module:** or **Version:**
        meta = re.match(r'^\*\*(.+?):\*\*\s+(.*)', line)
        if meta:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(meta.group(1) + ': ')
            r.font.size  = Pt(9.5)
            r.font.bold  = True
            r.font.color.rgb = GREEN_DARK
            r2 = p.add_run(meta.group(2))
            r2.font.size  = Pt(9.5)
            r2.font.color.rgb = GREY_LIGHT
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(3)
        add_inline(p, line, base_size=10)
        i += 1

# ── main ─────────────────────────────────────────────────────────────────────

def main():
    doc = setup_document()
    add_cover(doc)
    add_toc(doc)

    for idx, fname in enumerate(FILE_ORDER):
        path = os.path.join(DOCS_DIR, fname)
        if not os.path.exists(path):
            print(f"  [skip] {fname} not found")
            continue

        title = DOC_TITLES.get(fname, fname)
        print(f"  [{idx+1:02d}] {title}")

        with open(path, encoding='utf-8') as f:
            content = f.read()

        render_markdown(doc, content, title)

        if idx < len(FILE_ORDER) - 1:
            doc.add_page_break()

    doc.save(OUTPUT)
    print(f"\nSaved: {OUTPUT}")
    print(f"Pages: ~{len(FILE_ORDER) * 4}+  (open in Word for exact count)")

if __name__ == "__main__":
    main()
