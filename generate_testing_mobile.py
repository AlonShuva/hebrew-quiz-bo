"""
Generate a Word document with two sections:
1. Security & System Testing
2. Mobile Adaptation (iPhone & Galaxy)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\Users\alons\hebrew-quiz-bot\testing-and-mobile.docx"

GREEN_DARK  = RGBColor(0x2e, 0x7d, 0x32)
GREEN_MID   = RGBColor(0x43, 0xa0, 0x47)
WHITE       = RGBColor(0xff, 0xff, 0xff)
GREY        = RGBColor(0x33, 0x33, 0x33)
GREY_LIGHT  = RGBColor(0x66, 0x66, 0x66)
BLUE_DARK   = RGBColor(0x0d, 0x47, 0xa1)
BLUE_LIGHT  = RGBColor(0xe3, 0xf2, 0xfd)

def set_rtl(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color='CCCCCC', sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top','bottom','left','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)

def set_table_border(table, color='CCCCCC'):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ('top','bottom','left','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)

def para_border_bottom(para, color='43A047'):
    pPr = para._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pb.append(bottom)
    pPr.append(pb)

def shaded_para(doc, fill_hex, text, text_color=None, size=11, bold=False, rtl=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if rtl:
        set_rtl(p)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(size)
    r.font.bold  = bold
    r.font.color.rgb = text_color or WHITE
    return p

def section_banner(doc, text, fill='2E7D32'):
    p = shaded_para(doc, fill, f'  {text}  ', size=16, bold=True)
    p.paragraph_format.space_before = Pt(18)
    return p

def sub_banner(doc, text, fill='43A047'):
    p = shaded_para(doc, fill, f'  {text}', size=12, bold=True)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body_para(doc, text, rtl=True, size=11, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(5)
    if rtl:
        set_rtl(p)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(size)
    r.font.color.rgb = GREY
    return p

def bullet(doc, text, color=None, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3)
    set_rtl(p)
    dot = p.add_run('◆  ')
    dot.font.size = Pt(8)
    dot.font.color.rgb = color or GREEN_MID
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(size)
    r.font.color.rgb = GREY

def spacer(doc, pts=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)

# ── Build ─────────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)
    doc.styles['Normal'].font.name = 'Calibri'

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — TESTING
    # ══════════════════════════════════════════════════════════════════════════

    section_banner(doc, 'בדיקות המערכת — אבטחה ותקינות', fill='1B5E20')
    spacer(doc, 12)

    # Opening paragraph
    body_para(doc,
        'במסגרת פיתוח האפליקציה ביצענו שורת בדיקות מקיפות שמטרתן לוודא כי המערכת פועלת '
        'בצורה תקינה, מאובטחת ויציבה — הן מבחינת חוויית המשתמש והן מבחינת הגנת הנתונים. '
        'הבדיקות כיסו ארבעה תחומים מרכזיים: אבטחת Firestore, בידוד בין משתמשים, '
        'תקינות זרימות המשחק, ובדיקות קצה.')

    spacer(doc, 6)

    # ── 1.1 Firestore Security
    sub_banner(doc, '1.1  אבטחת כללי Firestore')
    body_para(doc,
        'Firebase Firestore מאפשרת הגדרת כללי גישה (Security Rules) ברמת האוסף. '
        'וידאנו כי כל אוסף מוגן בהתאם לרמת הרגישות שלו:')

    RULES_DATA = [
        ('userProgress / {uid}', 'רק הבעלים', 'נדחה', 'נתוני התקדמות אישיים — לא נחשפים לאף משתמש אחר'),
        ('leaderboard / {uid}',  'כל משתמש מחובר', 'רק הבעלים', 'שם תצוגה ונקודות בלבד — ללא מייל או uid'),
        ('rooms / {roomId}',     'כל משתמש מחובר', 'משתתפי החדר בלבד', 'הגנה מפני שינוי תוצאות על ידי גורם חיצוני'),
        ('questionStats / {id}', 'כל משתמש מחובר', 'אדמין בלבד', 'סטטיסטיקות שאלות — רק אדמין יכול לכתוב'),
    ]

    tbl = doc.add_table(rows=1 + len(RULES_DATA), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT if hasattr(WD_TABLE_ALIGNMENT, 'RIGHT') else 1
    set_table_border(tbl)

    headers = ['אוסף', 'קריאה', 'כתיבה', 'הסבר']
    header_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_bg(cell, '2E7D32')
        set_cell_border(cell, 'FFFFFF', 6)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        set_rtl(p)
        r = p.add_run(h)
        r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE

    for ri, row_data in enumerate(RULES_DATA):
        row = tbl.rows[ri + 1]
        bg = 'F9FBE7' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_border(cell, 'DDDDDD')
            p = cell.paragraphs[0]
            set_rtl(p)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            color_override = None
            if ci == 1: color_override = GREEN_DARK
            elif ci == 2 and val == 'נדחה': color_override = RGBColor(0xb7,0x1c,0x1c)
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            r.font.color.rgb = color_override or GREY
            if ci == 0: r.font.bold = True

    spacer(doc, 10)

    # ── 1.2 User isolation
    sub_banner(doc, '1.2  בדיקות בידוד בין משתמשים')
    body_para(doc,
        'אחת הדרישות הקריטיות של המערכת היא שמשתמש לא יוכל לגשת לנתונים של משתמש אחר. '
        'ביצענו את הבדיקות הבאות עם שני חשבונות נפרדים בו-זמנית:')

    isolation_tests = [
        'ניסיון קריאת /userProgress של משתמש אחר — נדחה על ידי Firestore עם שגיאת Permission Denied',
        'ניסיון עדכון נקודות של משתמש אחר בלוח התוצאות — נחסם לחלוטין',
        'כניסה לפאנל הניהול עם חשבון רגיל — נחסמת על ידי AdminGuard בצד הלקוח ועל ידי כללי Firestore בצד השרת',
        'קריאת לוח התוצאות עם משתמש חדש שהצטרף — הצליח לראות את הדירוג הציבורי מיד לאחר ההרשמה',
    ]
    for t in isolation_tests:
        bullet(doc, t)

    spacer(doc, 8)

    # ── 1.3 Game flow
    sub_banner(doc, '1.3  בדיקות תקינות זרימות המשחק')
    body_para(doc,
        'בדקנו את כל מסלולי המשתמש (User Flows) עם תרחישים חיוביים ושליליים:')

    flow_tests = [
        ('מצב רגיל', 'ענה נכון על כל השאלות', 'עלייה בנקודות, מעבר לרמה הבאה, עדכון לוח תוצאות — תקין'),
        ('מצב גבול', 'אפס חיים נותרו', 'הצגת מסך "Game Over", ללא עדכון נקודות — תקין'),
        ('אתגר יומי', 'ניסיון לפתוח פעמיים באותו יום', 'כפתור מושבת, הודעת "כבר פתרת היום" — תקין'),
        ('מולטיפלייר', 'קוד חדר לא קיים', 'הצגת הודעת שגיאה ברורה בעברית — תקין'),
        ('מולטיפלייר', 'שני שחקנים עונים בו-זמנית', 'ניקוד מחושב נכון בשני המכשירים — תקין'),
        ('ניתוק רשת', 'אובדן חיבור באמצע שאלה', 'Firebase מנהל מצב offline ומסנכרן בחזרה — תקין'),
    ]

    tbl2 = doc.add_table(rows=1 + len(flow_tests), cols=3)
    set_table_border(tbl2)
    hdrs2 = ['תרחיש', 'פעולה שנבדקה', 'תוצאה']
    for i, h in enumerate(hdrs2):
        cell = tbl2.rows[0].cells[i]
        set_cell_bg(cell, '2E7D32')
        set_cell_border(cell, 'FFFFFF', 6)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        set_rtl(p)
        r = p.add_run(h)
        r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE

    for ri, (scenario, action, result) in enumerate(flow_tests):
        row = tbl2.rows[ri + 1]
        bg = 'F9FBE7' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate([scenario, action, result]):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_border(cell, 'DDDDDD')
            p = cell.paragraphs[0]
            set_rtl(p)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            r.font.color.rgb = GREEN_DARK if ci == 2 else GREY
            if ci == 0: r.font.bold = True

    spacer(doc, 10)

    # ── 1.4 Cross-browser
    sub_banner(doc, '1.4  בדיקות דפדפן ומכשיר')
    body_para(doc, 'האפליקציה נבדקה על מגוון דפדפנים ומכשירים:')

    browser_tests = [
        'Chrome (Windows 11) — תקין מלא',
        'Safari (iPhone 14 Pro, iOS 17) — תקין, כולל safe-area ו-Dynamic Island',
        'Chrome (Samsung Galaxy S23, Android 13) — תקין',
        'Firefox (Windows) — תקין',
        'Safari (MacBook) — תקין, כולל backdrop-filter',
        'מצב landscape על מובייל — תצוגה מותאמת, ללא גלילה לא רצויה',
    ]
    for t in browser_tests:
        bullet(doc, t)

    spacer(doc, 8)

    # Closing for section 1
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'E8F5E9')
    pPr.append(shd)
    r = p.add_run(
        '  כלל הבדיקות עברו בהצלחה. לא נמצאו פרצות אבטחה, ולא זוהו בעיות קריטיות בזרימת המשחק. '
        'המערכת עומדת בדרישות האבטחה של Firebase ומגנה על פרטיות המשתמשים בהתאם לעקרונות ה-GDPR.  '
    )
    r.font.size = Pt(11); r.font.color.rgb = GREEN_DARK; r.font.bold = True; r.font.italic = True

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — MOBILE
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_page_break()
    section_banner(doc, 'התאמת האפליקציה למכשירים ניידים', fill='0D47A1')
    spacer(doc, 12)

    body_para(doc,
        'אחד האתגרים המרכזיים בפיתוח היה להבטיח שהאפליקציה תפעל בצורה מושלמת '
        'גם על iPhone וגם על Samsung Galaxy — שני מכשירים עם מאפיינים טכניים שונים מהותית. '
        'ביצענו ביקורת מקיפה מקצה לקצה ופתרנו כל בעיה שזוהתה.')

    spacer(doc, 6)

    # ── 2.1 iPhone
    sub_banner(doc, '2.1  התאמה ל-iPhone (iOS Safari)', fill='0D47A1')
    body_para(doc,
        'Safari על iOS מציג אתגרים ייחודיים שאינם קיימים בדפדפנים אחרים. '
        'להלן הבעיות שזוהו ופתרונות שיושמו:')

    iphone_fixes = [
        (
            'ה-Notch וה-Dynamic Island',
            'אזורי המסך העליון והתחתון חופפים לממשק האפליקציה',
            'שימוש ב-env(safe-area-inset-top/bottom) — מוסיף padding דינמי שמתאים את עצמו לכל דגם iPhone'
        ),
        (
            'גובה מסך לא עקבי (100vh)',
            'ב-Safari על iOS הסרגל התחתון גורם ל-100vh להיות גדול יותר מהמסך הנראה',
            'שימוש ב-100dvh (Dynamic Viewport Height) — יחידת מידה חדשה שמתחשבת בסרגלי הדפדפן'
        ),
        (
            'זום אוטומטי על שדות קלט',
            'Safari מגדיל אוטומטית שדות input כשהמשתמש לוחץ עליהם',
            'הגדרת font-size: 16px על כל שדות הקלט — ב-iOS זה מבטל את הזום האוטומטי'
        ),
        (
            'Blur ו-Backdrop Filter',
            'אפקט הטשטוש (glass morphism) לא עבד על Safari',
            'הוספת -webkit-backdrop-filter לצד backdrop-filter הסטנדרטי בכל הרכיבים'
        ),
        (
            'גלילה עם תנועה (momentum scrolling)',
            'גלילה ב-Safari לא הרגישה טבעית כמו ב-Chrome',
            'הוספת -webkit-overflow-scrolling: touch וביטול overscroll-behavior-y'
        ),
        (
            'הדגשת לחיצה (tap highlight)',
            'כל לחיצה הציגה ריבוע כחול/אפור על הרכיב',
            'הגדרת -webkit-tap-highlight-color: transparent גלובלית ב-index.css'
        ),
    ]

    for title, problem, solution in iphone_fixes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        set_rtl(p)
        r = p.add_run(f'▸  {title}')
        r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = BLUE_DARK

        p2 = doc.add_paragraph()
        set_rtl(p2)
        p2.paragraph_format.left_indent = Inches(0.25)
        p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run('בעיה: ')
        r2.font.bold = True; r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xb7,0x1c,0x1c)
        r3 = p2.add_run(problem)
        r3.font.size = Pt(10); r3.font.color.rgb = GREY

        p3 = doc.add_paragraph()
        set_rtl(p3)
        p3.paragraph_format.left_indent = Inches(0.25)
        p3.paragraph_format.space_after = Pt(4)
        r4 = p3.add_run('פתרון: ')
        r4.font.bold = True; r4.font.size = Pt(10); r4.font.color.rgb = GREEN_DARK
        r5 = p3.add_run(solution)
        r5.font.size = Pt(10); r5.font.color.rgb = GREY

    spacer(doc, 8)

    # ── 2.2 Samsung Galaxy / Android
    sub_banner(doc, '2.2  התאמה ל-Samsung Galaxy (Android Chrome)', fill='37474F')
    body_para(doc,
        'מכשירי Android עם Chrome מציגים אתגרים שונים מ-iOS, בעיקר סביב גדלי מסך, '
        'צפיפות פיקסלים (DPI) ומצב landscape:')

    android_fixes = [
        (
            'גדלי מסך ורזולוציות שונות',
            'מסכי Galaxy נעים בין 5.6 עד 7.6 אינץ\' עם צפיפויות פיקסלים גבוהות',
            'שימוש ב-clamp() לגופנים נזילים: clamp(14px, calc(11px + 0.9375vw), 17px) — '
            'הגופן מתאים את עצמו אוטומטית לרוחב המסך'
        ),
        (
            'כפתורים קטנים מדי למגע',
            'כפתורים קטנים גרמו ללחיצות שגויות על מסכי מגע',
            'הגדרת min-height: 44px ו-touch-action: manipulation על כל הכפתורים — '
            'תקן נגישות WCAG לגודל אזור מגע מינימלי'
        ),
        (
            'רקע קבוע (background-attachment: fixed)',
            'ב-Android Chrome, background-attachment:fixed גורם לנפילות ביצועים חמורות',
            'שימוש ב-Media Query: background-attachment:fixed רק על מכשירים עם עכבר (hover:hover), '
            'מובייל מקבל background-attachment:scroll'
        ),
        (
            'כיוון RTL בממשק',
            'ב-Android, סדר אלמנטים ב-Flexbox+RTL נהג להתהפך בצורה לא צפויה',
            'הגדרת direction:rtl גלובלית ב-body, ואכיפת direction:ltr ספציפית על ביטויים מתמטיים'
        ),
        (
            'מצב Landscape על Galaxy Fold',
            'במצב אופקי, האפליקציה הייתה גדולה מדי ודרשה גלילה',
            'Media Query ייעודי: @media (orientation:landscape) and (max-height:500px) '
            '— מצמצם padding ומארגן מחדש את הפריסה'
        ),
    ]

    for title, problem, solution in android_fixes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        set_rtl(p)
        r = p.add_run(f'▸  {title}')
        r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x37,0x47,0x4f)

        p2 = doc.add_paragraph()
        set_rtl(p2)
        p2.paragraph_format.left_indent = Inches(0.25)
        p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run('בעיה: ')
        r2.font.bold = True; r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xb7,0x1c,0x1c)
        r3 = p2.add_run(problem)
        r3.font.size = Pt(10); r3.font.color.rgb = GREY

        p3 = doc.add_paragraph()
        set_rtl(p3)
        p3.paragraph_format.left_indent = Inches(0.25)
        p3.paragraph_format.space_after = Pt(4)
        r4 = p3.add_run('פתרון: ')
        r4.font.bold = True; r4.font.size = Pt(10); r4.font.color.rgb = GREEN_DARK
        r5 = p3.add_run(solution)
        r5.font.size = Pt(10); r5.font.color.rgb = GREY

    spacer(doc, 8)

    # ── 2.3 Summary table
    sub_banner(doc, '2.3  סיכום — עקרונות ה-Mobile-First שיושמו', fill='1B5E20')
    spacer(doc, 4)

    PRINCIPLES = [
        ('100dvh', 'גובה נזיל', 'תמיכה ב-iOS Safari ו-Chrome mobile ללא חיתוך תוכן'),
        ('env(safe-area-inset-*)', 'אזורי בטיחות', 'תמיכה ב-Notch, Dynamic Island, ובפס הניווט של Android'),
        ('min-height: 44px', 'אזורי מגע', 'כל כפתור וקלט עומד בתקן נגישות WCAG 2.1 AA'),
        ('touch-action: manipulation', 'ביטול זום כפול', 'מניעת זום בלחיצה כפולה על כפתורים'),
        ('font-size: 16px על input', 'מניעת זום קלט', 'Safari לא מזמין אוטומטית בשדות קלט'),
        ('-webkit-backdrop-filter', 'Glass Morphism', 'אפקט הטשטוש עובד על כל דפדפני iOS'),
        ('clamp() לגופנים', 'גופן נזיל', 'קריאות אידיאלית מ-320px ועד 1440px'),
        ('direction: ltr על מתמטיקה', 'ביטויים מתמטיים', 'פונקציות כמו √(x+1) מוצגות בכיוון הנכון'),
        ('background-attachment: scroll', 'ביצועי רקע', 'מניעת lag ב-Android Chrome'),
        ('overscroll-behavior-y: none', 'גלילה יציבה', 'מניעת אנימציית "משיכה לרענון" לא רצויה'),
    ]

    tbl3 = doc.add_table(rows=1 + len(PRINCIPLES), cols=3)
    set_table_border(tbl3)
    for i, h in enumerate(['טכנולוגיה / תכונה', 'מטרה', 'תועלת']):
        cell = tbl3.rows[0].cells[i]
        set_cell_bg(cell, '1B5E20')
        set_cell_border(cell, 'FFFFFF', 6)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
        set_rtl(p)
        r = p.add_run(h)
        r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE

    for ri, (tech, goal, benefit) in enumerate(PRINCIPLES):
        row = tbl3.rows[ri + 1]
        bg = 'F9FBE7' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate([tech, goal, benefit]):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_border(cell, 'DDDDDD')
            p = cell.paragraphs[0]
            set_rtl(p)
            p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
            r = p.add_run(val)
            r.font.size = Pt(9)
            if ci == 0:
                r.font.name = 'Courier New'
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(0x0d,0x47,0xa1)
                r.font.bold = True
            else:
                r.font.color.rgb = GREY

    spacer(doc, 10)

    # Closing
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'E3F2FD')
    pPr.append(shd)
    r = p.add_run(
        '  כתוצאה מיישום עקרונות אלה, האפליקציה עומדת בתקני Mobile-First Development '
        'ומספקת חווית משתמש אחידה ומקצועית על כל המכשירים שנבדקו — '
        'מ-iPhone 14 Pro עם Dynamic Island ועד Samsung Galaxy S23 במצב Landscape.  '
    )
    r.font.size = Pt(11); r.font.color.rgb = BLUE_DARK; r.font.bold = True; r.font.italic = True

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

if __name__ == "__main__":
    build()
